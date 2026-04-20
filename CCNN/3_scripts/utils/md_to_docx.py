#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to Word (.docx) 转换脚本
将 EXPERIMENT_REPORT.md 转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading_with_style(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_from_markdown(doc, md_table_lines):
    """从 Markdown 表格生成 Word 表格"""
    if len(md_table_lines) < 2:
        return
    
    # 解析表头
    header_line = md_table_lines[0].strip('|').strip()
    headers = [h.strip() for h in header_line.split('|')]
    
    # 跳过分隔线
    # 解析数据行
    rows = []
    for line in md_table_lines[2:]:
        if line.strip():
            row_data = line.strip('|').strip()
            row = [cell.strip() for cell in row_data.split('|')]
            if len(row) == len(headers):
                rows.append(row)
    
    # 创建表格
    if rows:
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # 表头加粗
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 添加数据行
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_text in enumerate(row):
                row_cells[j].text = cell_text

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # 设置背景色
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E7E6E6')
    p._element.get_or_add_pPr().append(shading_elm)

def markdown_to_word(md_file, docx_file):
    """将 Markdown 文件转换为 Word 文档"""
    
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 标题处理
        if line.startswith('# '):
            text = line[2:].strip()
            add_heading_with_style(doc, text, 1)
            i += 1
        elif line.startswith('## '):
            text = line[3:].strip()
            add_heading_with_style(doc, text, 2)
            i += 1
        elif line.startswith('### '):
            text = line[4:].strip()
            add_heading_with_style(doc, text, 3)
            i += 1
        elif line.startswith('#### '):
            text = line[5:].strip()
            add_heading_with_style(doc, text, 4)
            i += 1
        
        # 表格处理
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, table_lines)
        
        # 代码块处理
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
            i += 1
        
        # 列表项处理（特殊标记）
        elif line.strip().startswith('- ') or line.strip().startswith('* ') or \
             (len(line.strip()) > 0 and line.strip()[0].isdigit() and '.' in line[:3]):
            p = doc.add_paragraph(line.strip()[2:].strip(), style='List Bullet')
            i += 1
        
        # 空行处理
        elif line.strip() == '':
            i += 1
        
        # 普通段落
        else:
            text = line.strip()
            if text:
                # 处理加粗、斜体
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 移除 ** 标记
                text = re.sub(r'__(.*?)__', r'\1', text)      # 移除 __ 标记
                text = re.sub(r'\*(.*?)\*', r'\1', text)      # 移除 * 标记
                text = re.sub(r'_(.*?)_', r'\1', text)        # 移除 _ 标记
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # 移除链接格式
                
                p = doc.add_paragraph(text)
            i += 1
    
    doc.save(docx_file)
    print(f"✓ 文档已生成: {docx_file}")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, '..', '..', '5_docs', 'reports', 'EXPERIMENT_REPORT.md')
    docx_file = os.path.join(script_dir, '..', '..', '5_docs', 'reports', 'EXPERIMENT_REPORT.docx')
    
    print("🔄 正在转换 Markdown 到 Word 格式...")
    markdown_to_word(md_file, docx_file)
    print("✓ 转换完成！")
